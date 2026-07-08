#!/usr/bin/env python
"""Build WSQ assessment documents for Microsoft Azure AI Fundamentals (AI-900).

Outputs are written to Assessments:
- WA (SAQ) question paper and answer guide
- PP Assessment question paper and answer guide

Question counts are fixed at 12 WA questions and 5 PP tasks.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Assessments"
ASSETS = ROOT / "courseware" / "assets"
TERTIARY_MARK = ASSETS / "imported-reference-media" / "dp900-image2.png"
COURSE_BADGE = ASSETS / "ai900-course-badge.png"

TITLE = "Microsoft Azure AI Fundamentals (AI-900)"
TITLE_FILE = "Microsoft Azure AI Fundamentals AI-900"
CODE = "TGS-2023021100"
ORG = "Tertiary Infotech Academy Pte Ltd"
UEN = "UEN: 201200696W"
URL = "www.tertiarycourses.com.sg"
VERSION = "v1"
DATE = "8 July 2026"

BLUE = RGBColor(0x1F, 0x6F, 0xEB)
INK = RGBColor(0x11, 0x18, 0x27)
GREY = RGBColor(0x55, 0x5B, 0x66)
LIGHT = "E8F0FE"


WRITTEN = [
    (
        "K1",
        "Slides 16, 24, 26; Lab 01",
        "A retailer wants to use AI for image checking, customer feedback analysis, invoice extraction, product recommendations and campaign copy. Explain how you would identify the AI workload type for each requirement.",
        [
            "Computer vision is used for image checking or OCR-style image/text extraction.",
            "Natural language processing or Azure AI Language is used for customer feedback analysis.",
            "Document Intelligence is used for invoice extraction and structured field capture.",
            "Machine learning can support prediction or recommendation scenarios.",
            "Generative AI supports drafting campaign copy or summarising content.",
        ],
    ),
    (
        "K2",
        "Slides 28, 72; Lab 01",
        "Describe the responsible AI principles that should be considered before deploying an Azure AI solution.",
        [
            "Fairness: check whether groups are affected unequally.",
            "Reliability and safety: test outputs and define fallback paths.",
            "Privacy and security: protect data, access, keys and logs.",
            "Inclusiveness: design for different users, languages and abilities.",
            "Transparency and accountability: explain AI use and name decision ownership.",
        ],
    ),
    (
        "K3",
        "Slides 29, 35-38; Lab 02",
        "Differentiate regression, classification and clustering using one workplace example for each.",
        [
            "Regression predicts a numeric value, such as future sales or delivery time.",
            "Classification predicts a category, such as fraud/not fraud or positive/negative feedback.",
            "Clustering groups similar records without pre-defined labels, such as customer segments.",
            "Good answers link the technique to the input data and expected output.",
        ],
    ),
    (
        "K4",
        "Slides 29, 39-42; Lab 03",
        "Explain the main stages in a machine learning lifecycle and how Azure Machine Learning supports them.",
        [
            "Define the prediction task and success measure.",
            "Prepare features and labels from trustworthy data.",
            "Train and evaluate a model using suitable metrics.",
            "Deploy, monitor and retrain the model over time.",
            "Azure Machine Learning supports experiment tracking, AutoML, model management, deployment and monitoring.",
        ],
    ),
    (
        "K5",
        "Slides 27, 43-58; Labs 04-07",
        "A business asks for image analysis, text sentiment analysis, speech transcription and document extraction. Map each requirement to the closest Azure AI service.",
        [
            "Image analysis maps to Azure AI Vision.",
            "Text sentiment or key phrase extraction maps to Azure AI Language.",
            "Speech transcription maps to Azure AI Speech.",
            "Translation maps to Azure AI Translator.",
            "Document extraction maps to Azure AI Document Intelligence.",
        ],
    ),
    (
        "K6",
        "Slides 43-46; Lab 04",
        "What should a learner check when validating output from a computer vision or OCR activity?",
        [
            "Whether image labels, detected objects or OCR text match the source image.",
            "Whether confidence scores are acceptable for the business use case.",
            "Whether errors need human review or correction.",
            "Whether the image contains personal or sensitive data that needs protection.",
        ],
    ),
    (
        "K7",
        "Slides 47-50; Lab 05",
        "Explain common natural language processing outputs and how they can support a business decision.",
        [
            "Sentiment analysis indicates positive, neutral or negative feedback.",
            "Key phrase extraction identifies important topics or terms.",
            "Named entity recognition identifies people, places, organisations or other entities.",
            "Language output can help route cases, summarise themes or prioritise review.",
        ],
    ),
    (
        "K8",
        "Slides 51-54; Lab 06",
        "Compare speech recognition, speech synthesis and translation in an Azure AI scenario.",
        [
            "Speech recognition converts spoken audio to text.",
            "Speech synthesis converts text to spoken audio.",
            "Translation converts text or speech from one language to another.",
            "A support assistant may combine all three for multilingual voice interaction.",
        ],
    ),
    (
        "K9",
        "Slides 55-58; Lab 07",
        "Explain how document intelligence and knowledge mining are different but related.",
        [
            "Document Intelligence extracts fields, tables and values from forms or documents.",
            "Knowledge mining or Azure AI Search indexes content so it can be searched and retrieved.",
            "Both help turn unstructured or semi-structured documents into usable information.",
            "Good answers mention validation of extracted fields and search relevance.",
        ],
    ),
    (
        "K10",
        "Slides 30, 59-62; Lab 08",
        "Describe the main controls needed when using generative AI, Azure AI Foundry or Azure OpenAI.",
        [
            "Use clear prompts and system instructions.",
            "Ground responses with source material where possible.",
            "Apply content filters and safety controls.",
            "Evaluate outputs for accuracy, bias and harmful content.",
            "Make sure users know when content is AI-assisted and keep human review for high-risk use.",
        ],
    ),
    (
        "K11",
        "Slides 63-66, 72; Lab 09",
        "List governance, privacy, security, cost and cleanup controls that should be included in an Azure AI solution design.",
        [
            "Use role-based access control and managed identities where appropriate.",
            "Protect personal or confidential data and avoid unnecessary data collection.",
            "Monitor logs, usage and quality.",
            "Use budgets, alerts and free tiers where possible.",
            "Clean up or stop resources after labs or tests.",
        ],
    ),
    (
        "K12",
        "Slides 67-74; Lab 10",
        "Given a new AI business scenario, what steps should you follow to recommend an Azure AI service and prepare for assessment review?",
        [
            "Identify the input type: image, text, speech, document, tabular data or prompt.",
            "Identify the required output: prediction, label, extracted field, transcript, translation, summary or answer.",
            "Choose the closest Azure AI service.",
            "Name privacy, security, human review and cost controls.",
            "Validate the result and clean up resources.",
        ],
    ),
]


SCENARIO = (
    "A retail and service organisation wants to introduce Azure AI capabilities across customer support, "
    "document processing, product quality review, multilingual communication, analytics and internal knowledge "
    "workflows. Complete the practical tasks below by following the sequence and evidence produced in the AI-900 labs."
)


PRACTICAL = [
    (
        "Task 1",
        "LO1",
        "Labs 01 and 10",
        "Classify AI workloads, map them to Azure AI services and create a responsible AI checklist for the organisation's scenario.",
        "Show your completed workload-service-control mapping and responsible AI checklist in the box below.",
        [
            "From Lab 01, list at least five business scenarios and classify each AI workload.",
            "Map each workload to the closest Azure service: Vision, Language, Speech, Document Intelligence, Azure Machine Learning, Azure AI Foundry or Azure OpenAI.",
            "Create a responsible AI checklist covering fairness, reliability, privacy/security, inclusiveness, transparency and accountability.",
            "From Lab 10, refine the service map into a final scenario-service-control recommendation.",
        ],
    ),
    (
        "Task 2",
        "LO2",
        "Labs 02 and 03",
        "Prepare a machine learning solution outline that identifies the ML technique, features, labels, lifecycle stages and Azure Machine Learning evidence.",
        "Show your ML technique table, lifecycle notes and Azure Machine Learning validation evidence in the box below.",
        [
            "From Lab 02, distinguish regression, classification and clustering for the given business examples.",
            "Identify possible features, labels and evaluation measures.",
            "From Lab 03, describe the Azure Machine Learning or AutoML lifecycle: prepare data, train, evaluate, deploy and monitor.",
            "State what evidence validates the model or workspace activity.",
        ],
    ),
    (
        "Task 3",
        "LO3",
        "Labs 04, 05, 06 and 07",
        "Complete an Azure AI services evidence pack for vision, language, speech/translation and document intelligence or search.",
        "Show screenshots or notes proving each service output and validation check in the box below.",
        [
            "From Lab 04, validate image analysis, OCR or object detection output and note any accuracy/privacy issue.",
            "From Lab 05, validate language outputs such as sentiment, key phrases or entities.",
            "From Lab 06, validate speech recognition, synthesis, translation or conversational AI output.",
            "From Lab 07, validate document extraction or knowledge mining/search evidence.",
            "For each service, state the input, output, Azure service used and human review point.",
        ],
    ),
    (
        "Task 4",
        "LO4",
        "Lab 08",
        "Use generative AI concepts to design a safe Azure AI Foundry or Azure OpenAI interaction for a business support scenario.",
        "Show your prompt, expected output, grounding/review notes and safety controls in the box below.",
        [
            "From Lab 08, define the user task and draft a clear prompt or system instruction.",
            "Identify whether the answer needs grounding or retrieved context.",
            "Name the relevant Azure AI Foundry, Azure OpenAI or model catalog concept used in class.",
            "Describe content safety, evaluation and human review controls.",
            "Record the final response quality check and any limitation.",
        ],
    ),
    (
        "Task 5",
        "LO5",
        "Labs 09 and 10",
        "Produce a final Azure AI solution design that includes service selection, governance, security, cost controls, validation and cleanup.",
        "Show your final design map, control checklist, validation evidence and cleanup plan in the box below.",
        [
            "From Lab 09, document requirements, selected service, data boundary, RBAC/access, privacy/security controls and cost considerations.",
            "Add monitoring, budget alert or free-tier considerations where applicable.",
            "From Lab 10, complete the capstone service map and course review evidence.",
            "State how learners validate the output and what resources must be deleted or stopped.",
            "Confirm the design aligns to the final service selection checklist from the slides.",
        ],
    ),
]


def set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, INK),
        ("Heading 3", 11.5, BLUE),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color


def para(doc: Document, text: str = "", size: float = 10.5, bold: bool = False, color: RGBColor | None = None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def cover(doc: Document, instrument: str) -> None:
    if TERTIARY_MARK.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(46)
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(str(TERTIARY_MARK), width=Inches(1.55))
    p = para(doc, ORG, 11.5, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(2)
    p = para(doc, UEN, 8.5, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(20)
    p = para(doc, instrument.upper(), 27, True, BLUE, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p = para(doc, "For", 9, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(8)
    if COURSE_BADGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(30)
        p.add_run().add_picture(str(COURSE_BADGE), width=Inches(1.35))
    p = para(doc, TITLE, 21, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(13)
    p = para(doc, f"TGS Ref No: {CODE}", 10, False, INK, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(12)
    p = para(doc, "Conducted by", 8.5, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(1)
    p = para(doc, ORG, 10.5, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(1)
    p = para(doc, UEN, 8.5, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(13)
    p = para(doc, f"Version {VERSION}", 11, True, BLUE, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def footer(doc: Document) -> None:
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = f"(c) 2026 {ORG}. All rights reserved. - {URL}"
    for run in p.runs:
        run.font.size = Pt(7.5)
        run.font.color.rgb = GREY


def metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
        set_cell_bg(cells[0], LIGHT)


def trainee_info(doc: Document) -> None:
    doc.add_paragraph("A. Trainee Information", style="Heading 1")
    metadata_table(doc, [
        ("Trainee Name", ""),
        ("Last 3 NRIC Digits + Alphabet", ""),
        ("Assessment Date", ""),
        ("Assessor Name", ""),
    ])


def official_use(doc: Document) -> None:
    doc.add_paragraph("For Official Use Only", style="Heading 1")
    metadata_table(doc, [
        ("Assessment Result", "Competent / Not Yet Competent"),
        ("Assessor Name", ""),
        ("Assessor NRIC", ""),
        ("Date", ""),
        ("Signature", ""),
    ])


def answer_box(doc: Document, height_pt: int = 95) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(height_pt * 20))
    tr_h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_h)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def build_written_question_doc() -> Path:
    doc = Document()
    style_doc(doc)
    cover(doc, "Written Assessment (SAQ)")
    para(doc, "Written Assessment (Short Answer Questions)", 16, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"{TITLE} - {CODE} - {VERSION}", 9, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    trainee_info(doc)
    doc.add_paragraph("B. Instructions to Candidate", style="Heading 1")
    add_bullets(doc, [
        "Answer all 12 open-ended questions.",
        "This assessment tests knowledge covered in the course slides and labs.",
        "Use approved course materials only as instructed by the assessor.",
        "Write clear, concise answers in the space provided.",
    ])
    doc.add_paragraph("C. Questions", style="Heading 1")
    for idx, (criterion, source, question, _answers) in enumerate(WRITTEN, 1):
        p = para(doc, f"Question {idx} ({criterion})", 11, True, BLUE)
        p.paragraph_format.keep_with_next = True
        para(doc, question)
        para(doc, f"Alignment: {source}", 8.5, False, GREY)
        answer_box(doc)
    official_use(doc)
    footer(doc)
    path = OUT / f"WA (SAQ) - {TITLE_FILE} - {VERSION}.docx"
    doc.save(path)
    return path


def build_written_answer_doc() -> Path:
    doc = Document()
    style_doc(doc)
    cover(doc, "Answer to Written Assessment (SAQ)")
    para(doc, "Answer Guide - Written Assessment (SAQ)", 16, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"{TITLE} - {CODE} - {VERSION}", 9, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("Marking Guidance", style="Heading 1")
    add_bullets(doc, [
        "Award marks where the candidate covers the idea, even if wording differs.",
        "Each question is aligned to the final course slides and labs.",
        "Do not require examples outside the AI-900 courseware.",
    ])
    for idx, (criterion, source, question, answers) in enumerate(WRITTEN, 1):
        doc.add_paragraph(f"Question {idx} ({criterion})", style="Heading 2")
        para(doc, question)
        para(doc, f"Source alignment: {source}", 8.5, False, GREY)
        para(doc, "Suggestive answers (not exhaustive):", 10.5, True)
        add_bullets(doc, answers)
    footer(doc)
    path = OUT / f"Answer to WA (SAQ) - {TITLE_FILE} - {VERSION}.docx"
    doc.save(path)
    return path


def build_pp_question_doc() -> Path:
    doc = Document()
    style_doc(doc)
    cover(doc, "Practical Performance Assessment")
    para(doc, "Practical Performance Assessment", 16, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"{TITLE} - {CODE} - {VERSION}", 9, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    trainee_info(doc)
    doc.add_paragraph("B. Instructions to Candidate", style="Heading 1")
    add_bullets(doc, [
        "Complete all 5 practical tasks.",
        "The practical performance must follow the AI-900 labs completed in class.",
        "Use screenshots, completed tables, notes or output evidence from the labs.",
        "Do not submit personal or confidential data.",
    ])
    doc.add_paragraph("C. Scenario", style="Heading 1")
    para(doc, SCENARIO)
    doc.add_paragraph("D. Practical Tasks", style="Heading 1")
    for label, criterion, labs, prompt, caption, _points in PRACTICAL:
        doc.add_paragraph(f"{label} ({criterion}) - {labs}", style="Heading 2")
        para(doc, prompt)
        para(doc, caption, 9, True, GREY)
        answer_box(doc, 150)
    official_use(doc)
    footer(doc)
    path = OUT / f"PP Assessment - {TITLE_FILE} - {VERSION}.docx"
    doc.save(path)
    return path


def build_pp_answer_doc() -> Path:
    doc = Document()
    style_doc(doc)
    cover(doc, "Answer to Practical Performance Assessment")
    para(doc, "Answer Guide - Practical Performance Assessment", 16, True, INK, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"{TITLE} - {CODE} - {VERSION}", 9, False, GREY, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("Scenario", style="Heading 1")
    para(doc, SCENARIO)
    doc.add_paragraph("Marking Guidance", style="Heading 1")
    add_bullets(doc, [
        "Award competence when the learner provides valid lab-aligned evidence and can explain the service choice.",
        "PP evidence may include screenshots, completed tables, validation notes, control checklists and cleanup notes.",
        "The assessor should check the practical evidence against the listed lab sources.",
    ])
    for label, criterion, labs, prompt, _caption, points in PRACTICAL:
        doc.add_paragraph(f"{label} ({criterion}) - {labs}", style="Heading 2")
        para(doc, prompt)
        para(doc, "Expected practical evidence / model performance points:", 10.5, True)
        add_bullets(doc, points)
    footer(doc)
    path = OUT / f"Answer to PP Assessment - {TITLE_FILE} - {VERSION}.docx"
    doc.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    paths = [
        build_written_question_doc(),
        build_written_answer_doc(),
        build_pp_question_doc(),
        build_pp_answer_doc(),
    ]
    for path in paths:
        print(f"Generated {path}")
    print(f"WA questions: {len(WRITTEN)}")
    print(f"PP tasks: {len(PRACTICAL)}")


if __name__ == "__main__":
    main()
